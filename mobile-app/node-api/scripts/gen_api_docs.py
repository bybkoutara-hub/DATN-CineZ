# -*- coding: utf-8 -*-
"""Sinh tài liệu danh sách API (Word .docx) cho backend node-api của app mobile."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

YELLOW = RGBColor(0xE2, 0xA4, 0x3B)
DARK = RGBColor(0x15, 0x15, 0x17)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Danh sách API: nhóm, method, đường dẫn, mô tả, body/params, auth, màn hình app, trạng thái
APIS = [
    ("Xác thực (Auth)", "POST", "/api/auth/register",
     "Đăng ký tài khoản mới", "name, email, password, phone", "Không",
     "(Signup giữ OTP)", "Có hàm, chưa gắn UI"),
    ("Xác thực (Auth)", "POST", "/api/auth/login",
     "Đăng nhập, trả về JWT token (30 ngày)", "email, password", "Không",
     "sign-in.tsx", "Đã tích hợp"),
    ("Xác thực (Auth)", "GET", "/api/auth/profile",
     "Lấy thông tin cá nhân theo token", "—", "Bearer JWT",
     "profile.tsx", "Đã tích hợp"),
    ("Phim (Movies)", "GET", "/api/movies",
     "Danh sách phim (lọc theo ?status=now_playing|coming_soon)", "query: status", "Không",
     "index.tsx, movie.tsx", "Đã tích hợp"),
    ("Phim (Movies)", "POST", "/api/movies",
     "Thêm phim mới (chức năng quản trị)", "title, poster_url, duration, genres, status, release_date", "Không",
     "—", "Không áp dụng (admin)"),
    ("Phim (Movies)", "GET", "/api/movies/:id",
     "Chi tiết phim kèm danh sách suất chiếu", "param: id", "Không",
     "movie-detail.tsx", "Đã tích hợp"),
    ("Suất chiếu (Showtimes)", "POST", "/api/movies/showtimes",
     "Thêm suất chiếu (chức năng quản trị)", "movieId, roomName, startTime, price", "Không",
     "—", "Không áp dụng (admin)"),
    ("Suất chiếu (Showtimes)", "GET", "/api/movies/showtimes/:id",
     "Chi tiết 1 suất chiếu + danh sách ghế trống", "param: id", "Không",
     "select-seat.tsx", "Đã tích hợp"),
    ("Đặt vé (Bookings)", "POST", "/api/bookings",
     "Tạo đơn đặt vé (giữ ghế, trừ ghế trống, lưu combo)", "showtimeId, seats[], combos[], totalPrice, paymentMethod", "Bearer JWT",
     "payment.tsx", "Đã tích hợp"),
    ("Đặt vé (Bookings)", "GET", "/api/bookings/my-history",
     "Lịch sử vé đã đặt của người dùng hiện tại", "—", "Bearer JWT",
     "ticket.tsx", "Đã tích hợp"),
]

STATUS_COLOR = {
    "Đã tích hợp": RGBColor(0x1E, 0x7E, 0x34),
    "Có hàm, chưa gắn UI": RGBColor(0xB8, 0x86, 0x00),
    "Không áp dụng (admin)": RGBColor(0x6C, 0x75, 0x7D),
}


def set_cell_bg(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# Tiêu đề
title = doc.add_heading('', level=0)
run = title.add_run('TÀI LIỆU DANH SÁCH API')
run.font.color.rgb = YELLOW
run.font.size = Pt(24)
run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Hệ thống đặt vé xem phim — Backend node-api (App Mobile)')
r.italic = True
r.font.size = Pt(12)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Base URL: http://192.168.1.30:5001/api  •  Express + TypeScript + MongoDB + JWT').font.size = Pt(9)

doc.add_paragraph()

# Tổng quan
doc.add_heading('1. Tổng quan', level=1)
total = len(APIS)
done = sum(1 for a in APIS if a[7] == 'Đã tích hợp')
admin = sum(1 for a in APIS if 'admin' in a[7])
p = doc.add_paragraph()
p.add_run(f'• Tổng số endpoint: {total}\n')
p.add_run(f'• Đã tích hợp vào app mobile: {done}\n')
p.add_run(f'• Chức năng quản trị (app khách không dùng): {admin}\n')
p.add_run('• Còn lại: POST /auth/register (giữ luồng OTP nên chưa gắn nút UI).')

# Bảng chi tiết
doc.add_heading('2. Bảng chi tiết các API', level=1)
headers = ['Method', 'Endpoint', 'Mô tả', 'Tham số / Body', 'Auth', 'Màn hình app', 'Trạng thái']
table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].text = ''
    run = hdr[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9.5)
    set_cell_bg(hdr[i], '1F1F23')

for grp, method, path, desc, body, auth, screen, status in APIS:
    cells = table.add_row().cells
    vals = [method, path, desc, body, auth, screen, status]
    for i, v in enumerate(vals):
        cells[i].text = ''
        run = cells[i].paragraphs[0].add_run(v)
        run.font.size = Pt(9)
        if i == 0:  # method
            run.bold = True
            run.font.color.rgb = YELLOW
        if i == 6 and status in STATUS_COLOR:  # trạng thái
            run.font.color.rgb = STATUS_COLOR[status]
            run.bold = True

# Nhóm theo chức năng
doc.add_heading('3. Phân nhóm theo chức năng', level=1)
groups = {}
for grp, method, path, *_ in APIS:
    groups.setdefault(grp, []).append(f'{method} {path}')
for grp, items in groups.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{grp}: ').bold = True
    p.add_run(', '.join(items)).font.size = Pt(9.5)

# Ghi chú
doc.add_heading('4. Ghi chú', level=1)
notes = [
    "Tất cả endpoint cần đăng nhập phải gửi kèm header: Authorization: Bearer <token>.",
    "Token JWT được lưu ở AsyncStorage sau khi đăng nhập và tự động đính kèm qua axios interceptor.",
    "API POST /bookings đã được bổ sung cơ chế fallback: chạy không transaction khi MongoDB ở chế độ standalone.",
    "Phản hồi thống nhất dạng JSON: { success, message, data } (hoặc { success, data }).",
]
for n in notes:
    doc.add_paragraph(n, style='List Number')

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run('Tài liệu sinh tự động bằng Python (python-docx)')
fr.italic = True
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

OUT = 'DANH_SACH_API.docx'
doc.save(OUT)
print('Da tao file:', OUT)
